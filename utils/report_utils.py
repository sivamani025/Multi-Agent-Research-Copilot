import os

from docx import Document

from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer
)

from reportlab.lib.styles import (
    getSampleStyleSheet
)


def generate_docx(
    topic,
    summary,
    critique,
    metrics
):
    """
    Generate DOCX report.
    """

    os.makedirs(
        "reports/docx",
        exist_ok=True
    )

    safe_topic = (
        topic.replace("/", "_")
        .replace("\\", "_")
    )

    filepath = (
        f"reports/docx/{safe_topic}.docx"
    )

    doc = Document()

    # Title

    doc.add_heading(
        topic,
        level=1
    )

    # Summary

    doc.add_heading(
        "Executive Summary",
        level=2
    )

    doc.add_paragraph(
        summary
    )

    # Critique

    doc.add_heading(
        "Critical Analysis",
        level=2
    )

    doc.add_paragraph(
        critique
    )

    # Metrics

    doc.add_heading(
        "Evaluation Metrics",
        level=2
    )

    for key, value in metrics.items():

        doc.add_paragraph(
            f"{key}: {value}"
        )

    doc.save(
        filepath
    )

    return filepath


def generate_pdf(
    topic,
    summary,
    critique,
    metrics
):
    """
    Generate PDF report.
    """

    os.makedirs(
        "reports/pdf",
        exist_ok=True
    )

    safe_topic = (
        topic.replace("/", "_")
        .replace("\\", "_")
    )

    filepath = (
        f"reports/pdf/{safe_topic}.pdf"
    )

    pdf = SimpleDocTemplate(
        filepath
    )

    styles = getSampleStyleSheet()

    content = []

    # ----------------------------------
    # Title
    # ----------------------------------

    content.append(
        Paragraph(
            topic,
            styles["Title"]
        )
    )

    content.append(
        Spacer(1, 20)
    )

    # ----------------------------------
    # Summary
    # ----------------------------------

    content.append(
        Paragraph(
            "Executive Summary",
            styles["Heading2"]
        )
    )

    content.append(
        Paragraph(
            summary,
            styles["BodyText"]
        )
    )

    content.append(
        Spacer(1, 12)
    )

    # ----------------------------------
    # Critique
    # ----------------------------------

    content.append(
        Paragraph(
            "Critical Analysis",
            styles["Heading2"]
        )
    )

    content.append(
        Paragraph(
            critique,
            styles["BodyText"]
        )
    )

    content.append(
        Spacer(1, 12)
    )

    # ----------------------------------
    # Metrics
    # ----------------------------------

    content.append(
        Paragraph(
            "Evaluation Metrics",
            styles["Heading2"]
        )
    )

    for key, value in metrics.items():

        content.append(
            Paragraph(
                f"{key}: {value}",
                styles["BodyText"]
            )
        )

    content.append(
        Spacer(1, 12)
    )

    pdf.build(
        content
    )

    return filepath